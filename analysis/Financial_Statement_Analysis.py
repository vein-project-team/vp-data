# -*- coding: utf-8 -*-

import pandas as pd
import numpy as np
from tqdm import tqdm as pb
import datetime
import re
import warnings

import matplotlib.pyplot as plt
import pylab as mpl

from docx import Document
from docx.shared import Pt

from data_source import local_source

def concat_ts_codes(df): #拼接df中所有TS_CODE为输入条件的格式
    result = ''
    for code in df["TS_CODE"]:
        result = result + 'TS_CODE = "' + code + '" or ' 
    result = result[:-4]
    return result


def drop_duplicates_keep_nonnan(df,subset): #保留nan最少的行, 暂时没用
    warnings.filterwarnings("ignore")
    subset_values = []
    df_result = pd.DataFrame(columns=df.columns)
    for i in range(len(df)):
        subset_value = list(df[subset].iloc[i,:])
        if subset_value not in subset_values: subset_values.append(subset_value)
    for subset_value in subset_values:
        df_sub = df[(df[subset]==subset_value).product(axis=1)==1]
        df_sub["nan_count"] = 0
        df_sub.loc[:,"nan_count"] = df_sub.isnull().sum(axis=1)
        df_sub.sort_values(by='nan_count',ascending=True, inplace=True)
        df_sub = pd.DataFrame(df_sub.iloc[0,:]).T
        df_result = pd.concat([df_result, df_sub],axis=0)    
    warnings.filterwarnings("default")
    return df_result


#tester = pd.DataFrame([[1,1,5,5,5],[1,1,5,np.nan,np.nan],[2,2,5,5,5],[2,2,np.nan,5,5],[2,1,np.nan,np.nan,np.nan]],columns=['a','b','c','d','e'])
#tester2 = drop_duplicates_keep_nonnan(df=tester, subset=['a','b'])


def Find_Comparibles(ts_code, df_ind): #按总市值差选取对比公司, 可改进。输入df需要有END_DATE和INDUSTRY和TOTAL_MV列
    stocks_used = df_ind.copy()
    
    stocks_used["END_DATE"] = stocks_used["END_DATE"].astype(int)
    last_end_date = max(stocks_used["END_DATE"])
    stocks_used = stocks_used[stocks_used["END_DATE"]==last_end_date]
    
    stocks_used["TOTAL_MV_diff"] = abs( stocks_used["TOTAL_MV"] - stocks_used.loc[stocks_used["TS_CODE"]==ts_code, "TOTAL_MV"].iloc[0] )
    stocks_used.sort_values(by="TOTAL_MV_diff", ascending=True,inplace=True)
    stocks_used = stocks_used[1:3]
    return list(stocks_used["TS_CODE"])


def RatioComparation_Plotter(df,var_name, year=5):
    fig = plt.figure(figsize=(4, 4))
    ax = fig.add_subplot(1,1,1)        
    ax.set_title("{var_name}趋势图".format(var_name=var_name))                 
    ax.set_xlabel("年份",labelpad=0, position=(0.5,1))           
    ax.set_ylabel("{var_name}".format(var_name=var_name),labelpad=0, position=(1,0.5))
    
    for stock in df["TS_CODE"].unique():
        x = df.loc[df["TS_CODE"]==stock,"END_DATE_year"].iloc[(-1*year):]
        y = df.loc[df["TS_CODE"]==stock, var_name].iloc[(-1*year):]
        ax.plot(x,y,linewidth='1',label="{stock}".format(stock=stock))
    ax.legend(loc="upper right",bbox_to_anchor=(1.4,1),shadow=True)
    plt.show()


def FSA_Initializer(ts_code):
    basic = local_source.get_stock_list(condition='TS_CODE = '+'"'+ts_code+'"')
    ind = basic["INDUSTRY"].iloc[0]
    stocks_ind = local_source.get_stock_list(condition='INDUSTRY = '+'"'+ind+'"')
    ts_codes_ind = concat_ts_codes(stocks_ind) 
    quotations_monthly_ind = local_source.get_quotations_monthly(cols='TRADE_DATE,TS_CODE,CLOSE',condition=ts_codes_ind).sort_values(by="TRADE_DATE", ascending=True) 
    quotations_monthly_ind.rename(columns={'TRADE_DATE':'END_DATE'}, inplace = True)
    stock_indicators_daily_ind = local_source.get_stock_indicators_daily(cols='TRADE_DATE,TS_CODE,TOTAL_SHARE',condition=ts_codes_ind).sort_values(by="TRADE_DATE", ascending=True)
    stock_indicators_daily_ind.rename(columns={'TRADE_DATE':'END_DATE'}, inplace = True)
    financial_indicators_ind = local_source.get_financial_indicators(condition=ts_codes_ind).sort_values(by="END_DATE", ascending=True)
    stocks_ind = pd.merge(financial_indicators_ind,stocks_ind, on=['TS_CODE'], how="left") 
    stocks_ind = pd.merge(stocks_ind, quotations_monthly_ind, on=['TS_CODE','END_DATE'], how="left") 
    stocks_ind = pd.merge(stocks_ind, stock_indicators_daily_ind, on=['TS_CODE','END_DATE'], how="left") 
    stocks_ind = stocks_ind.applymap(lambda x: np.nan if x=="NULL" else x)
    stocks_ind["TOTAL_MV"] = stocks_ind["TOTAL_SHARE"] * stocks_ind["CLOSE"]
    stocks_ind["END_DATE_year"]=[datetime.datetime.strptime(str(date), "%Y%m%d").year for date in stocks_ind["END_DATE"]]
    stocks_ind["END_DATE_month"]=[datetime.datetime.strptime(str(date), "%Y%m%d").month for date in stocks_ind["END_DATE"]]
    stocks_ind = stocks_ind[stocks_ind["END_DATE_month"]==12]
    stocks_ind_avg=stocks_ind.groupby(["END_DATE"]).mean().reset_index()
    stocks_ind_avg["TS_CODE"]='AVERAGE'
    stocks_used = Find_Comparibles(ts_code, stocks_ind)
    stocks_used.append(ts_code)
    stocks_ind_used = [i for i in stocks_ind.index if stocks_ind.loc[i,"TS_CODE"] in stocks_used]
    stocks_ind_used = stocks_ind.loc[stocks_ind_used,:]
    stocks_ind_used = pd.concat([stocks_ind_used, stocks_ind_avg],axis=0).reset_index(drop=True)
    return stocks_ind_used


def FSA_Analyzer(df): #该df应由FSA_Initializer输出
    mpl.rcParams['font.sans-serif'] = ['FangSong'] 
    mpl.rcParams['axes.unicode_minus'] = False
    mpl.rcParams['font.weight'] = 'bold'
    var_name_list = ['QUICK_RATIO','CURRENT_RATIO','CASH_RATIO','DEBT_TO_EQT','AR_TURN','ASSETS_TURN','ROA','ROE','EBIT_TO_INTEREST']
    for var_name in var_name_list:
        RatioComparation_Plotter(df=df, var_name=var_name)



if __name__ == '__main__':
    ts_code = '000002.SZ'
    stocks_ind_used = FSA_Initializer(ts_code=ts_code)
    FSA_Analyzer(stocks_ind_used)



'''
liquidity: quick ratio, currency ratio
solvency: debt to equity ratio
operating capability: AR turnover, asset turnover
profitability: ROA, ROE, EBITDA/Interest expense, cash from operations to total debt
'''



'''
document = Document()
paragraph = document.add_paragraph('Nor的mal text, ')
paragraph.add_run('add text')
style = document.styles['Normal']
font = style.font
font.size = Pt(10)
document.save('test.docx')
'''
